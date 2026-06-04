"""Generate test proof documents in DOCX format.

Transforms text log files from automated tests into formatted Word documents,
including screenshots.

Architecture:
    - Recursive reading of log tree (campaigns > suites > cases)
    - UTC metadata parsing and local time conversion
    - Automatic screenshot detection and insertion
    - Robust error handling with logging

Output format:
    - Heading 1: Test campaign
    - Heading 2: Test suite
    - Heading 3: Test case
    - Body: Logs with inserted screenshots

Example:
    >>> generate_docx_proof(
    ...     logs_root=Path("logs/e2e"),
    ...     docx_root=Path(".docx_test_proofs_xxx"),
    ...     logger=logger,
    ... )

"""

import os
import re
import uuid
from concurrent.futures import ThreadPoolExecutor
from contextlib import suppress
from datetime import datetime
from functools import partial
from pathlib import Path
from typing import TYPE_CHECKING, NamedTuple, TypedDict, final

from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches

if TYPE_CHECKING:
    from collections.abc import Callable, Iterator

    from docx.document import Document as DocxDocument

    from ocarina.ports.ilogger import ILogger

_DEFAULT_SCREENSHOT_NEEDLE = "Screenshot: "
_DEFAULT_UTC_DATE_REGEX = re.compile(r"\[UTC_DATE::([^]]+)]")


def _default_format_date(local_dt: datetime) -> str:
    """Render a parsed UTC marker as local time: ``[MM/DD/YYYY | HHhMM:SS.ffffff]``.

    Receives the marker's datetime already converted to the local timezone, and
    returns the full text that replaces the ``[UTC_DATE::...]`` marker (brackets
    included). Override it via ``generate_docx_proof(format_date=...)`` to render
    another layout, e.g. a French ``[%d/%m/%Y à %Hh%M:%S]``.
    """
    return local_dt.strftime("[%m/%d/%Y | %Hh%M:%S.%f]")


def _replace_utc_date(
    line: str,
    *,
    utc_date_regex: re.Pattern[str],
    format_date: Callable[[datetime], str],
) -> str:
    def _repl(m: re.Match[str]) -> str:
        with suppress(Exception):
            local_dt = datetime.fromisoformat(
                m.group(1).replace("Z", "+00:00")  # noqa: FURB162
            ).astimezone()
            return format_date(local_dt)
        return m.group(0)

    return utc_date_regex.sub(_repl, line)


def _shorten_docx_path(docx_path: Path) -> Path:
    """Atomically reserve a short unique path if the stem exceeds 8 chars.

    When shortening is required, the candidate name is created with
    ``O_CREAT | O_EXCL``, so the filesystem itself guarantees uniqueness in a
    single reserve-or-fail syscall (exactly like ``mkdir(exist_ok=False)``) —
    no check-then-act race between two concurrent workers. The reserved file is
    an empty placeholder the caller is expected to overwrite.

    When the stem already fits, the path is returned unchanged and nothing is
    reserved.

    Raises:
        RuntimeError: If a unique filename cannot be reserved after 500 attempts.

    """
    max_stem_length = 8
    retries = 500

    if len(docx_path.stem) <= max_stem_length:
        return docx_path

    parent = docx_path.parent
    for _ in range(retries):
        short_name = uuid.uuid4().hex[:max_stem_length]
        new_path = parent / f"{short_name}.docx"
        try:
            fd = os.open(new_path, os.O_CREAT | os.O_EXCL | os.O_WRONLY, 0o644)
        except FileExistsError:
            continue
        os.close(fd)
        return new_path

    msg = (  # pragma: no cover
        f"Cannot generate a unique short filename for:"
        " "
        f"{docx_path} after {retries} attempts."
    )
    raise RuntimeError(msg)  # pragma: no cover


def _create_unique_output_dir(output_root: Path) -> Path:
    """Create a uniquely named subdirectory (4-char hex) inside output_root.

    Raises:
        RuntimeError: If a unique directory cannot be created after 500 attempts.

    """
    retries = 500

    for _ in range(retries):
        candidate = output_root / uuid.uuid4().hex[:4]
        try:
            candidate.mkdir(parents=True, exist_ok=False)
        except FileExistsError:
            continue
        else:
            return candidate

    msg = (
        "Cannot generate a unique output subdirectory under:"
        " "
        f"{output_root} after {retries} attempts."
    )
    raise RuntimeError(msg)


class _TestCaseEntry(TypedDict):
    test_campaign_name: str
    test_suite_name: str
    test_case_name: str
    file_path: Path


class _GenerationStats(NamedTuple):
    """Outcome of a generation run.

    ``attempted`` counts the test cases discovered in the log tree;
    ``succeeded`` counts those whose DOCX was actually written to disk.
    """

    attempted: int
    succeeded: int


def _safe_iter_lines(file_path: Path, logger: ILogger) -> Iterator[str]:
    try:
        with file_path.open("r", encoding="utf-8", errors="replace") as f:
            yield from f
    except Exception as exc:
        msg = f"Cannot read: {file_path}"
        logger.exception(msg, exc=exc)


def _save_docx(*, doc: DocxDocument, docx_path: Path, logger: ILogger) -> bool:
    try:
        doc.save(str(docx_path))
    except Exception as exc:
        msg = f"Cannot save file: {docx_path}"
        logger.exception(msg, exc=exc)
        return False
    else:
        logger.success(f"Successfully written: {docx_path}")
        return True


@final
class _DocxProofGenerator:
    """Generate DOCX test proofs from log files.

    Expected log tree:
        logs_root/
        ├── campaign_1/
        │   ├── suite_A/
        │   │   ├── test_case_1.log
        │   │   └── test_case_2.log
        │   └── suite_B/
        │       └── test_case_3.log
        └── campaign_2/
            └── ...
    """

    def __init__(
        self,
        *,
        logs_root: Path,
        docx_root: Path,
        screenshot_needle: str = _DEFAULT_SCREENSHOT_NEEDLE,
        utc_date_regex: re.Pattern[str] = _DEFAULT_UTC_DATE_REGEX,
        format_date: Callable[[datetime], str] = _default_format_date,
    ) -> None:
        self.logs_root = logs_root
        self.docx_root = docx_root
        self._screenshot_needle = screenshot_needle
        self._utc_date_regex = utc_date_regex
        self._format_date = format_date

    def _iter_test_cases(self, logger: ILogger) -> Iterator[_TestCaseEntry]:
        try:
            for campaign_dir in self.logs_root.iterdir():
                if not campaign_dir.is_dir():
                    continue

                for suite_dir in campaign_dir.iterdir():
                    if not suite_dir.is_dir():
                        continue

                    for file in suite_dir.iterdir():
                        if file.is_file():
                            yield _TestCaseEntry(
                                test_campaign_name=campaign_dir.name,
                                test_suite_name=suite_dir.name,
                                test_case_name=file.stem,
                                file_path=file,
                            )
        except OSError as exc:
            msg = f"Failed to iterate over: {self.logs_root}"
            logger.exception(msg, exc=exc)

    def _create_docx_from_case(
        self, test_case: _TestCaseEntry, logger: ILogger
    ) -> bool:
        campaign_name = test_case["test_campaign_name"]
        suite_name = test_case["test_suite_name"]
        case_name = test_case["test_case_name"]
        file_path = test_case["file_path"]

        doc = Document()
        doc.add_heading(
            campaign_name, level=1
        ).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_heading(suite_name, level=2).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        doc.add_paragraph()
        doc.add_heading(case_name, level=3)

        for line in _safe_iter_lines(file_path, logger):
            normalized_line = _replace_utc_date(
                line.rstrip("\n"),
                utc_date_regex=self._utc_date_regex,
                format_date=self._format_date,
            )

            if self._screenshot_needle in normalized_line:
                parts = normalized_line.split(self._screenshot_needle, 1)
                if len(parts) > 1 and parts[1].strip():
                    image_path = Path(parts[1].strip())
                    if image_path.exists() and image_path.is_file():
                        doc.add_paragraph()
                        try:
                            doc.add_picture(str(image_path), width=Inches(6))
                        except Exception as exc:
                            msg = f"Cannot add image {image_path}"
                            logger.exception(msg, exc=exc)
            else:
                p = doc.add_paragraph()
                p.paragraph_format.space_before = 0
                p.paragraph_format.space_after = 0
                p.add_run(normalized_line)

        relative_path = file_path.relative_to(self.logs_root)
        docx_path = (self.docx_root / relative_path).with_suffix(".docx").resolve()

        try:
            docx_path.parent.mkdir(parents=True, exist_ok=True)
            doc.save(str(docx_path))
        except OSError:
            original_docx_path = docx_path
            docx_path = _shorten_docx_path(docx_path)
            msg = (
                f"Filename too long or access error for: {original_docx_path}."
                f" Retrying with: {docx_path}"
            )
            logger.warning(msg)
            saved = _save_docx(docx_path=docx_path, doc=doc, logger=logger)
            if not saved:
                # Drop the empty placeholder reserved above (or any partial
                # file) so a failed retry never leaves a 0-byte .docx behind.
                with suppress(OSError):
                    docx_path.unlink()
            return saved
        else:
            return True

    def generate_docx_proofs(
        self, logger: ILogger, *, max_workers: int = 1
    ) -> _GenerationStats:
        """Generate one DOCX per test case, optionally in parallel.

        Cases are independent units of disk I/O (each reads its own log, builds
        its own Document, and writes a unique output path), so they parallelise
        cleanly. The only shared object is ``logger``; its writes may interleave
        across threads, which is harmless for this best-effort reporter.

        Args:
            logger: Logger for progress and error reporting.
            max_workers: Upper bound on worker threads. ``<= 1`` runs the original
                sequential path verbatim — no list materialisation, no pool, no
                thread. Above 1, it is clamped to at most the number of cases.

        Returns:
            A ``_GenerationStats`` pair: the number of cases discovered
            (``attempted``) and the number whose DOCX was actually written
            to disk (``succeeded``).

        """
        if max_workers <= 1:
            attempted = 0
            succeeded = 0
            for case in self._iter_test_cases(logger):
                attempted += 1
                succeeded += self._create_docx_from_case(case, logger)
            return _GenerationStats(attempted=attempted, succeeded=succeeded)

        cases = list(self._iter_test_cases(logger))
        if not cases:
            return _GenerationStats(attempted=0, succeeded=0)

        workers = min(max_workers, len(cases))
        create = partial(self._create_docx_from_case, logger=logger)
        with ThreadPoolExecutor(max_workers=workers) as executor:
            succeeded = sum(executor.map(create, cases))
        return _GenerationStats(attempted=len(cases), succeeded=succeeded)


def generate_docx_proof(  # noqa: PLR0913
    *,
    logs_root: Path,
    output_root: Path,
    logger: ILogger,
    screenshot_needle: str = _DEFAULT_SCREENSHOT_NEEDLE,
    utc_date_regex: re.Pattern[str] = _DEFAULT_UTC_DATE_REGEX,
    format_date: Callable[[datetime], str] = _default_format_date,
    auto_create_unique_directory: bool = True,
    max_workers: int = 1,
) -> None:
    """Generate DOCX test proofs from log files.

    Args:
        logs_root: Root directory of the log tree to process.
        output_root: Root directory where DOCX files will be written.
        logger: Logger for progress and error reporting.
        screenshot_needle: used to detect screenshot lines. Default: "Screenshot: ".
        utc_date_regex: used to detect and replace UTC dates. Default: [UTC_DATE::...].
        format_date: Renders each matched date marker. Receives the marker's
            datetime already converted to local time and returns the full
            replacement text (brackets included). Default: a US-style
            ``[MM/DD/YYYY | HHhMM:SS.ffffff]``. Pass your own, e.g.
            ``lambda dt: dt.strftime("[%d/%m/%Y à %Hh%M:%S]")`` for a French layout.
        auto_create_unique_directory: creates automatically a random-named unique dir.
        max_workers: Worker threads used to generate documents in parallel.
            Default: 1 (sequential). Clamped to at least 1 and at most the total
            number of documents to generate; raise it to parallelise.

    """
    if not logs_root.is_dir():
        msg = f"Directory not accessible: {logs_root}. Generation skipped."
        logger.warning(msg)
        return

    docx_root = (
        _create_unique_output_dir(output_root)
        if auto_create_unique_directory
        else output_root
    )

    stats = _DocxProofGenerator(
        logs_root=logs_root,
        docx_root=docx_root,
        screenshot_needle=screenshot_needle,
        utc_date_regex=utc_date_regex,
        format_date=format_date,
    ).generate_docx_proofs(logger, max_workers=max_workers)

    if stats.attempted == 0:
        msg = f"No test case found under: {logs_root}. Nothing was generated."
        logger.warning(msg)
        if auto_create_unique_directory:
            with suppress(Exception):
                docx_root.rmdir()
        return

    if stats.succeeded == 0:
        msg = (
            f"Found {stats.attempted} test case(s) under: {logs_root}, but every"
            f" DOCX generation failed. Output: {docx_root}"
        )
        logger.warning(msg)
        return

    if stats.succeeded < stats.attempted:
        msg = (
            f"Plugin execution done with errors. Generated"
            f" {stats.succeeded}/{stats.attempted} DOCX. Output: {docx_root}"
        )
        logger.warning(msg)
        return

    msg = (
        f"Plugin execution done. Generated {stats.succeeded} DOCX. Output: {docx_root}"
    )
    logger.info(msg)
