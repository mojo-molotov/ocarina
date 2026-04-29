"""DOCX test-proof generator.

Covers:
- End-to-end: log tree → DOCX tree with the expected hierarchy and content.
- Screenshot lines: present file → picture inserted; missing file → line ignored.
- UTC date markers replaced with a human-readable local time.
- No-case / missing root: warning path + cleanup of auto-created dirs.
- auto_create_unique_directory=False uses output_root directly.
- Non-directory entries in the log tree are skipped without crashing.
- OSError from doc.save triggers the shortened-path fallback.
"""

# ruff: noqa: S101

import os
import re
import struct
import sys
import zlib
from typing import TYPE_CHECKING
from unittest.mock import patch

import allure
import pytest
from docx import Document
from docx.document import Document as DocxDocument

from ocarina.opinionated.plugins.reports.docx_tests_proofs import generate_docx_proof

if TYPE_CHECKING:
    from pathlib import Path

EPIC = "DOCX test-proofs"
FEATURE = "DOCX proofs"
LAYER = "integration"


class RecordingLogger:
    """Captures every logger call. Mirrors the subset docx uses."""

    def __init__(self) -> None:  # noqa: D107
        self.info_msgs: list[str] = []
        self.warning_msgs: list[str] = []
        self.success_msgs: list[str] = []
        self.exception_msgs: list[tuple[str, Exception | None]] = []

    def info(self, msg: str, *args, **kwargs) -> None:  # noqa: ARG002, D102
        self.info_msgs.append(msg)

    def warning(self, msg: str, *args, **kwargs) -> None:  # noqa: ARG002, D102
        self.warning_msgs.append(msg)

    def success(self, msg: str, *args, **kwargs) -> None:  # noqa: ARG002, D102
        self.success_msgs.append(msg)

    def exception(  # noqa: D102
        self,
        msg: str,
        *args,  # noqa: ARG002
        exc: Exception | None = None,
        **kwargs,  # noqa: ARG002
    ) -> None:
        self.exception_msgs.append((msg, exc))


def _minimal_png() -> bytes:
    """Build a valid 1x1 RGB PNG so python-docx can read it back."""
    sig = b"\x89PNG\r\n\x1a\n"

    def chunk(kind: bytes, data: bytes) -> bytes:
        crc = zlib.crc32(kind + data)
        return struct.pack(">I", len(data)) + kind + data + struct.pack(">I", crc)

    ihdr = struct.pack(">IIBBBBB", 1, 1, 8, 2, 0, 0, 0)
    idat = zlib.compress(b"\x00" + b"\xff\xff\xff")
    return sig + chunk(b"IHDR", ihdr) + chunk(b"IDAT", idat) + chunk(b"IEND", b"")


def _write_log(root: Path, campaign: str, suite: str, case: str, body: str) -> Path:
    suite_dir = root / campaign / suite
    suite_dir.mkdir(parents=True, exist_ok=True)
    log_path = suite_dir / f"{case}.log"
    log_path.write_text(body)
    return log_path


def _only_docx(tree: Path) -> list[Path]:
    return sorted(p for p in tree.rglob("*.docx") if p.is_file())


def _headings(doc_path: Path) -> list[tuple[str, int]]:
    doc = Document(str(doc_path))
    result: list[tuple[str, int]] = []
    for p in doc.paragraphs:
        style = p.style.name  # type: ignore[union-attr]
        if style.startswith("Heading "):
            result.append((p.text, int(style.rsplit(" ", 1)[-1])))
    return result


def _body_texts(doc_path: Path) -> list[str]:
    doc = Document(str(doc_path))
    return [
        p.text
        for p in doc.paragraphs
        if not p.style.name.startswith("Heading") and p.text  # type: ignore[union-attr]
    ]


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "happy-path", "hierarchy")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.CRITICAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "Single log file produces a DOCX with the expected campaign/suite/case heading hierarchy"  # noqa: E501
)
def test_single_case_generates_structured_docx(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    _write_log(logs_root, "main", "login", "case1", "ℹ️ opened\n✅ done\n")  # noqa: RUF001
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 1
    assert _headings(produced[0]) == [("main", 1), ("login", 2), ("case1", 3)]
    bodies = _body_texts(produced[0])
    assert any("opened" in b for b in bodies)
    assert any("done" in b for b in bodies)
    assert any("Plugin execution done" in m for m in logger.info_msgs)


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "happy-path", "multi-case")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("Every test case in every suite of every campaign produces its own DOCX")  # type: ignore[no-untyped-call,untyped-decorator]
def test_multiple_cases_produce_one_docx_each(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    _write_log(logs_root, "main", "login", "happy", "ok")
    _write_log(logs_root, "main", "login", "unhappy", "fail")
    _write_log(logs_root, "main", "upload", "happy", "ok")
    _write_log(logs_root, "smoke", "ping", "case", "ok")
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 4  # noqa: PLR2004


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "screenshot")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("A 'Screenshot: <path>' line pointing at a real image inserts a picture")  # type: ignore[no-untyped-call,untyped-decorator]
def test_existing_screenshot_is_inserted(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    img = tmp_path / "shot.png"
    img.write_bytes(_minimal_png())
    body = f"before\nScreenshot: {img}\nafter\n"
    _write_log(logs_root, "main", "suite", "case", body)
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 1
    doc = Document(str(produced[0]))
    assert len(doc.inline_shapes) == 1
    bodies = _body_texts(produced[0])
    assert any("before" in b for b in bodies)
    assert any("after" in b for b in bodies)
    # The "Screenshot: ..." line itself is replaced by the picture.
    assert not any("Screenshot:" in b for b in bodies)


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "screenshot", "error")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "A 'Screenshot: <path>' line pointing at a missing file produces no picture and no crash"  # noqa: E501
)
def test_missing_screenshot_is_silently_skipped(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    body = "before\nScreenshot: /nowhere/does-not-exist.png\nafter\n"
    _write_log(logs_root, "main", "suite", "case", body)
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 1
    doc = Document(str(produced[0]))
    assert len(doc.inline_shapes) == 0
    # The broken line is simply dropped, not logged as an exception.
    assert logger.exception_msgs == []


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "utc-date")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "[UTC_DATE::...] markers are replaced with a local formatted date in the DOCX body"
)
def test_utc_date_marker_is_formatted_locally(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    body = "[UTC_DATE::2025-03-15T12:34:56.789Z] hello\n"
    _write_log(logs_root, "main", "suite", "case", body)
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    bodies = _body_texts(produced[0])
    assert not any("[UTC_DATE::" in b for b in bodies)
    # Format is [MM/DD/YYYY | HHhMM:SS.ffffff]
    formatted_re = re.compile(r"\[\d{2}/\d{2}/\d{4} \| \d{2}h\d{2}:\d{2}\.\d+\]")
    assert any(formatted_re.search(b) for b in bodies)


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "utc-date", "robustness")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.MINOR)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("Malformed [UTC_DATE::...] markers are left untouched, not swallowed")  # type: ignore[no-untyped-call,untyped-decorator]
def test_malformed_utc_date_is_left_untouched(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    body = "[UTC_DATE::not-a-date] line\n"
    _write_log(logs_root, "main", "suite", "case", body)
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    bodies = _body_texts(_only_docx(output_root)[0])
    assert any("[UTC_DATE::not-a-date]" in b for b in bodies)


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "empty", "error")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.MINOR)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("Missing logs_root yields a warning and no DOCX")  # type: ignore[no-untyped-call,untyped-decorator]
def test_missing_logs_root_warns_and_does_nothing(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "does-not-exist"
    output_root = tmp_path / "out"
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    assert any("Directory not accessible" in m for m in logger.warning_msgs)
    assert not output_root.exists() or _only_docx(output_root) == []


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "empty", "cleanup")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "Empty logs_root produces a 'no test case found' warning and cleans up the auto-created dir"  # noqa: E501
)
def test_empty_logs_root_cleans_up_its_unique_dir(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    logs_root.mkdir()
    output_root = tmp_path / "out"
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    assert any("No test case found" in m for m in logger.warning_msgs)
    # The 4-char unique subdir is removed on the empty path.
    assert list(output_root.iterdir()) == []


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "config")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("auto_create_unique_directory=False writes directly under output_root")  # type: ignore[no-untyped-call,untyped-decorator]
def test_no_auto_unique_directory_writes_directly(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    _write_log(logs_root, "main", "suite", "case", "body")
    logger = RecordingLogger()

    generate_docx_proof(
        logs_root=logs_root,
        output_root=output_root,
        logger=logger,  # type: ignore[arg-type]
        auto_create_unique_directory=False,
    )

    # DOCX lives at output_root/main/suite/case.docx — no hex-named parent dir.
    assert (output_root / "main" / "suite" / "case.docx").exists()
    # And there is no additional 4-char hex dir next to it.
    top_children = [p for p in output_root.iterdir() if p.is_dir()]
    assert top_children == [output_root / "main"]


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "robustness")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title("Stray files at any level of the log tree are ignored without crashing")  # type: ignore[no-untyped-call,untyped-decorator]
def test_non_directory_entries_in_log_tree_are_skipped(tmp_path: Path) -> None:  # noqa: D103
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    _write_log(logs_root, "main", "suite", "case", "body")
    (logs_root / "loose_file.txt").write_text("noise")
    (logs_root / "main" / "loose_suite.txt").write_text("noise")
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 1
    assert logger.exception_msgs == []


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "fallback", "posix")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "A filename beyond the filesystem limit triggers the shortened-path fallback"
)
@pytest.mark.skipif(
    sys.platform == "win32" or not hasattr(os, "pathconf"),
    reason=(
        "triggers rely on POSIX NAME_MAX enforcement;"
        " Windows uses a path-total limit that interacts with pytest's tmp_path"
        " and can't be probed portably"
    ),
)
def test_too_long_filename_triggers_shortened_path_retry(tmp_path: Path) -> None:  # noqa: D103
    # Probe NAME_MAX dynamically so the test adapts to whatever the filesystem
    # actually enforces (255 on ext4/APFS/HFS+, larger on some exotic FSes).
    # Pick a stem such that ".log" (source) fits and ".docx" (target) overflows.
    name_max = os.pathconf(str(tmp_path), "PC_NAME_MAX")  # type: ignore[attr-defined,unused-ignore]
    stem_len = name_max - len(".log")  # e.g. 251 on typical POSIX
    assert stem_len + len(".docx") > name_max  # sanity

    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    long_case = "x" * stem_len
    _write_log(logs_root, "main", "suite", long_case, "body")
    logger = RecordingLogger()

    generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    produced = _only_docx(output_root)
    assert len(produced) == 1
    # Shortened filename → stem is 8-char hex, not the 251-char original.
    assert len(produced[0].stem) == 8  # noqa: PLR2004
    assert any("Filename too long" in m for m in logger.warning_msgs)
    assert any("Successfully written" in m for m in logger.success_msgs)


@allure.epic(EPIC)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.feature(FEATURE)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.tag("docx", "fallback", "portable")  # type: ignore[no-untyped-call,untyped-decorator]
@allure.severity(allure.severity_level.NORMAL)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.label("layer", LAYER)  # type: ignore[no-untyped-call,untyped-decorator]
@allure.title(  # type: ignore[no-untyped-call,untyped-decorator]
    "Same fallback — portable version: an OSError from doc.save triggers the shortened-path retry on any OS"  # noqa: E501
)
def test_oserror_from_save_triggers_shortened_path_retry_portable(  # noqa: D103
    tmp_path: Path,
) -> None:
    # Cross-platform twin of the test above. Instead of relying on the
    # filesystem to enforce NAME_MAX, we simulate the OSError by patching
    # python-docx's save method once, then letting the real save run on the
    # shortened path. Runs everywhere, including Windows where probing
    # NAME_MAX is not possible.
    logs_root = tmp_path / "logs"
    output_root = tmp_path / "out"
    long_case = "x" * 40  # just needs to exceed the 8-char shortening threshold
    _write_log(logs_root, "main", "suite", long_case, "body")
    logger = RecordingLogger()

    calls = {"n": 0}
    real_save = DocxDocument.save

    def flaky_save(self, path, *args, **kwargs):  # noqa: ANN001, ANN202
        calls["n"] += 1
        if calls["n"] == 1:
            msg = "simulated ENAMETOOLONG"
            raise OSError(msg)
        return real_save(self, path, *args, **kwargs)

    with patch.object(DocxDocument, "save", flaky_save):
        generate_docx_proof(logs_root=logs_root, output_root=output_root, logger=logger)  # type: ignore[arg-type]

    assert calls["n"] == 2  # one failing call + one successful retry  # noqa: PLR2004
    produced = _only_docx(output_root)
    assert len(produced) == 1
    assert len(produced[0].stem) == 8  # noqa: PLR2004
    assert any("Filename too long" in m for m in logger.warning_msgs)
    assert any("Successfully written" in m for m in logger.success_msgs)
